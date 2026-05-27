from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), kwargs.get('val', 'single'))
        border.set(qn('w:sz'), kwargs.get('sz', '6'))
        border.set(qn('w:color'), kwargs.get('color', '4472C4'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

doc = Document()

# ---------- 页面边距 ----------
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ---------- 标题 ----------
title = doc.add_heading('独立站各推广渠道考核方案', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.font.size = Pt(22)
    run.font.bold = True

# ---------- 副标题 ----------
sub = doc.add_paragraph('制定部门：市场营销部　　文件版本：V1.0　　制定日期：2026年4月16日')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()

# ========== 一、概述 ==========
doc.add_heading('一、方案概述', level=1)
doc.add_paragraph(
    '本方案旨在建立科学、量化的独立站多渠道推广考核体系，覆盖 SEO、'
    'SEM/Google Ads、社交媒体广告（Facebook/Instagram/TikTok）、'
    'KOL/网红营销、EDM 邮件营销及联盟营销六大渠道，'
    '通过统一的 KPI 框架与分级奖惩机制，提升整体 ROI，'
    '支撑独立站年度 GMV 目标达成。'
)

# ========== 二、考核总览 ==========
doc.add_heading('二、考核指标总览', level=1)

overview_data = [
    ['渠道', '核心 KPI', '权重占比', '目标基准'],
    ['SEO 自然搜索', '自然流量 / 关键词排名 / 转化率', '20%', '月均自然访客 ≥ 5 万'],
    ['SEM / Google Ads', 'ROAS / CPC / CTR / 转化成本', '25%', 'ROAS ≥ 3.5'],
    ['社交媒体广告', 'CPM / CPC / CTR / 购买转化率', '20%', 'CTR ≥ 1.5%'],
    ['KOL / 网红营销', '曝光量 / 互动率 / 带货 GMV', '15%', '单次合作 ROI ≥ 2'],
    ['EDM 邮件营销', '打开率 / 点击率 / 营收贡献', '10%', '打开率 ≥ 22%'],
    ['联盟营销', '订单量 / 联盟佣金率 / 新客占比', '10%', '月新增联盟订单 ≥ 300'],
]

tbl = doc.add_table(rows=len(overview_data), cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

header_color = '1F497D'
row_colors   = ['D9E2F3', 'FFFFFF']

for r_idx, row_data in enumerate(overview_data):
    row = tbl.rows[r_idx]
    for c_idx, text in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        if r_idx == 0:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            set_cell_bg(cell, header_color)
        else:
            run.font.size = Pt(9.5)
            bg = row_colors[(r_idx - 1) % 2]
            set_cell_bg(cell, bg)

doc.add_paragraph()

# ========== 三、各渠道详细考核方案 ==========
doc.add_heading('三、各渠道详细考核方案', level=1)

# --- 3.1 SEO ---
doc.add_heading('3.1 SEO 自然搜索', level=2)
doc.add_paragraph(
    'SEO 考核分为流量、排名与转化三个维度，每月对比上月数据，'
    '并以季度为周期进行综合复盘。'
)

seo_data = [
    ['考核指标', '计算方式', '目标值', '达成奖励'],
    ['自然搜索访客数', '月度 GA 数据', '≥ 50,000 UV', '达标奖金 ×1.0'],
    ['核心关键词 Top3 数量', '每月 SEMrush 导出', '≥ 20 个', '每增加 5 个 +5%'],
    ['自然流量转化率', 'GA 目标转化 / 自然会话', '≥ 2.5%', '超出 0.5% 加权 +10%'],
    ['页面平均加载速度', 'PageSpeed Insights', '≤ 2.5s', '未达标扣 5%'],
]
tbl2 = doc.add_table(rows=len(seo_data), cols=4)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for r_idx, row_data in enumerate(seo_data):
    row = tbl2.rows[r_idx]
    for c_idx, text in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        if r_idx == 0:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            set_cell_bg(cell, '2E74B5')
        else:
            run.font.size = Pt(9.5)
            set_cell_bg(cell, row_colors[(r_idx - 1) % 2])

doc.add_paragraph()

# --- 3.2 SEM ---
doc.add_heading('3.2 SEM / Google Ads', level=2)
doc.add_paragraph(
    'Google Ads 以 ROAS 为核心考核指标，辅以 CPC、CTR 和转化成本，'
    '月度结算，季度绩效系数核定。'
)

sem_data = [
    ['考核指标', '计算方式', '目标值', '达成奖励'],
    ['ROAS（广告回报率）', '广告带来营收 / 广告花费', '≥ 3.5', '超出 0.5 加权 +8%'],
    ['CPC（每次点击成本）', '总花费 / 总点击', '≤ $0.80', '低于目标 10% +5%'],
    ['CTR（点击率）', '点击 / 展示', '≥ 5%', '超出 1% 加权 +5%'],
    ['CPA（每次转化成本）', '总花费 / 转化数', '≤ $18', '低于目标 20% +10%'],
]
tbl3 = doc.add_table(rows=len(sem_data), cols=4)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
for r_idx, row_data in enumerate(sem_data):
    row = tbl3.rows[r_idx]
    for c_idx, text in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        if r_idx == 0:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            set_cell_bg(cell, '2E74B5')
        else:
            run.font.size = Pt(9.5)
            set_cell_bg(cell, row_colors[(r_idx - 1) % 2])

doc.add_paragraph()

# --- 3.3 社交媒体广告 ---
doc.add_heading('3.3 社交媒体广告（Facebook / Instagram / TikTok）', level=2)
doc.add_paragraph(
    '社交媒体广告按平台分别设置考核基准，统一上报汇总后计算综合得分。'
)

social_data = [
    ['平台', 'CPM 目标', 'CTR 目标', '购买转化率目标', 'ROAS 目标'],
    ['Facebook',  '≤ $8',   '≥ 1.5%', '≥ 2.0%', '≥ 3.0'],
    ['Instagram', '≤ $9',   '≥ 1.2%', '≥ 1.8%', '≥ 2.8'],
    ['TikTok',    '≤ $6',   '≥ 2.0%', '≥ 1.5%', '≥ 2.5'],
]
tbl4 = doc.add_table(rows=len(social_data), cols=5)
tbl4.style = 'Table Grid'
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
for r_idx, row_data in enumerate(social_data):
    row = tbl4.rows[r_idx]
    for c_idx, text in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        if r_idx == 0:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            set_cell_bg(cell, '2E74B5')
        else:
            run.font.size = Pt(9.5)
            set_cell_bg(cell, row_colors[(r_idx - 1) % 2])

doc.add_paragraph()

# --- 3.4 KOL ---
doc.add_heading('3.4 KOL / 网红营销', level=2)
doc.add_paragraph(
    'KOL 合作采用"保底 + 佣金"模式，考核以单次合作 ROI 和品牌曝光为主。'
)
kol_items = [
    '曝光量（Impression）：单次合作 ≥ 50 万次',
    '互动率（Engagement Rate）：≥ 3%（点赞 + 评论 + 分享 / 粉丝数）',
    '带货 GMV：单次合作 ≥ $5,000',
    'ROI：带货 GMV / 总合作费用 ≥ 2.0',
    '新客占比：通过 KOL 链接引流的新客 ≥ 60%',
]
for item in kol_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(10)

doc.add_paragraph()

# --- 3.5 EDM ---
doc.add_heading('3.5 EDM 邮件营销', level=2)
doc.add_paragraph(
    'EDM 按月度发送效果及营收贡献进行双维度考核，重点优化列表质量与内容转化。'
)
edm_data = [
    ['考核指标', '目标值', '说明'],
    ['邮件打开率',  '≥ 22%',   '行业均值约 18%'],
    ['点击率 (CTR)', '≥ 3.5%', '点击 / 发送总量'],
    ['退订率',      '≤ 0.3%',  '超出扣除当月绩效系数 0.1'],
    ['邮件带来营收', '≥ $15,000/月', '通过 UTM 追踪'],
    ['列表增长率',   '≥ 5%/月',      '净增有效订阅人数'],
]
tbl5 = doc.add_table(rows=len(edm_data), cols=3)
tbl5.style = 'Table Grid'
tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
for r_idx, row_data in enumerate(edm_data):
    row = tbl5.rows[r_idx]
    for c_idx, text in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        if r_idx == 0:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            set_cell_bg(cell, '2E74B5')
        else:
            run.font.size = Pt(9.5)
            set_cell_bg(cell, row_colors[(r_idx - 1) % 2])

doc.add_paragraph()

# --- 3.6 联盟营销 ---
doc.add_heading('3.6 联盟营销（Affiliate Marketing）', level=2)
doc.add_paragraph(
    '联盟营销通过第三方联盟平台管理，以订单量、GMV 及联盟健康度为考核维度。'
)
aff_items = [
    '月联盟订单量：≥ 300 单',
    '联盟 GMV：≥ $30,000 / 月',
    '新客占比：≥ 70%',
    '联盟商数量（活跃）：≥ 50 个',
    '平均佣金率：控制在 8%–12% 之间',
]
for item in aff_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(10)

doc.add_paragraph()

# ========== 四、绩效评级与奖惩 ==========
doc.add_heading('四、绩效评级与奖惩机制', level=1)

rating_data = [
    ['评级', '综合完成率', '绩效系数', '奖惩说明'],
    ['S（卓越）', '≥ 120%', '×1.3', '额外季度奖金 +20%'],
    ['A（优秀）', '100%–119%', '×1.1', '正常发放 + 优秀激励'],
    ['B（良好）', '85%–99%',  '×1.0', '正常发放'],
    ['C（待改进）','70%–84%', '×0.85','绩效谈话 + 改进计划'],
    ['D（不达标）','< 70%',   '×0.7', '末位警告 / 绩效辅导'],
]
tbl6 = doc.add_table(rows=len(rating_data), cols=4)
tbl6.style = 'Table Grid'
tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
grade_colors = {'S':'1E7B34','A':'2E74B5','B':'2E74B5','C':'C55A11','D':'C00000'}
for r_idx, row_data in enumerate(rating_data):
    row = tbl6.rows[r_idx]
    for c_idx, text in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        if r_idx == 0:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            set_cell_bg(cell, '1F497D')
        else:
            run.font.size = Pt(9.5)
            set_cell_bg(cell, row_colors[(r_idx - 1) % 2])

doc.add_paragraph()

# ========== 五、数据报告与复盘 ==========
doc.add_heading('五、数据报告与复盘周期', level=1)
review_items = [
    '日报：SEM / 社交媒体广告核心指标（花费、转化、ROAS），次日 10:00 前提交。',
    '周报：各渠道数据汇总 + 异常分析，每周一 12:00 前提交。',
    '月报：完整 KPI 达成情况 + 下月优化方向，每月 3 日前提交。',
    '季度复盘：渠道预算重新分配 + 年度目标校准，每季末最后一周完成。',
]
for item in review_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(10)

doc.add_paragraph()

# ========== 六、附则 ==========
doc.add_heading('六、附则', level=1)
doc.add_paragraph(
    '本方案自 2026 年 5 月 1 日起执行，由市场营销部负责解释与修订，'
    '如遇市场重大变化（如汇率波动 ≥10%、平台政策调整等），'
    '可经管理层审批后临时调整考核基准，调整结果在当月月报中注明。'
)

# ---------- 保存 ----------
output_path = r'D:\3857\Desktop\claude_code_projects\tesk_day_2\独立站各推广渠道的考核方案.docx'
doc.save(output_path)
print(f'文件已生成：{output_path}')
